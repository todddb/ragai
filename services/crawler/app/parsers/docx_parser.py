from __future__ import annotations

import io

from docx import Document

from app.parsers.types import ParsedDocument


def parse_docx(content_bytes: bytes) -> ParsedDocument:
    document = Document(io.BytesIO(content_bytes))
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    title = document.core_properties.title or (paragraphs[0] if paragraphs else "")
    markdown = "\n\n".join(paragraphs)
    text = " ".join(paragraphs)
    return ParsedDocument(
        title=title,
        markdown=markdown,
        text_for_chunking=text,
        meta={"paragraph_count": len(paragraphs)},
    )
